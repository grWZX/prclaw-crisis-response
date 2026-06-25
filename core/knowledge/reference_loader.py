#!/usr/bin/env python3
"""
Reference sources loader for PR RAG/Graph/RLHF.

- 加载公关方法论文档
- 加载案例库结构化表（渠道、案例、目标、行业-品牌关系、关系表详情）
- 提供别名索引、schema 扩展提示和实体/关系校验辅助
"""

from __future__ import annotations

import json
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

import pandas as pd

try:  # 可选依赖
    import docx
except Exception:  # pragma: no cover - docx 不一定安装
    docx = None


DEFAULT_BASE_DIR = Path(os.environ.get("REFERENCE_SOURCES_DIR", "data/reference"))

DEFAULT_FILES = {
    "methodology": "公关营销传播方法论.md",
    "channels": "公关案例库_传播渠道关系表_表格.csv",
    "cases": "公关案例库_公关案例库_表格.csv",
    "goals": "公关案例库_公关目标关系表_表格.csv",
    "industry_brand": "公关案例库_行业与品牌关系表_表格.csv",
    "relation_details": "关系表详情.docx",
}


@dataclass
class SchemaExtension:
    node_types: Set[str] = field(default_factory=set)
    relationship_types: Set[str] = field(default_factory=set)
    predicates: Set[str] = field(default_factory=set)

    def to_prompt(self) -> str:
        nodes = ", ".join(sorted(self.node_types)) or "无"
        rels = ", ".join(sorted(self.relationship_types)) or "无"
        preds = ", ".join(sorted(self.predicates)) or "无"
        return (
            f"额外节点: {nodes}\n"
            f"额外关系: {rels}\n"
            f"可用谓词: {preds}"
        )


class ReferenceSources:
    """Load and cache local reference artifacts to guide NER/RE/RAG."""

    def __init__(self, base_dir: Path | str | None = None) -> None:
        self.base_dir = Path(base_dir) if base_dir is not None else DEFAULT_BASE_DIR
        self.files = DEFAULT_FILES
        self._methodology_text: Optional[str] = None
        self._case_tables: Dict[str, pd.DataFrame] = {}
        self._alias_index: Dict[str, str] = {}
        self._schema_extension: Optional[SchemaExtension] = None

    # ------------------------
    # Public API
    # ------------------------
    def methodology_text(self) -> str:
        """Return raw text of 公关营销传播方法论.md (empty string if missing)."""
        if self._methodology_text is not None:
            return self._methodology_text
        path = self.base_dir / self.files["methodology"]
        if not path.exists():
            self._methodology_text = ""
            return self._methodology_text
        self._methodology_text = path.read_text(encoding="utf-8", errors="ignore")
        return self._methodology_text

    def case_tables(self) -> Dict[str, pd.DataFrame]:
        """Load all CSV tables into memory (cached)."""
        if self._case_tables:
            return self._case_tables

        def _load_csv(key: str) -> pd.DataFrame:
            csv_path = self.base_dir / self.files[key]
            if not csv_path.exists():
                return pd.DataFrame()
            try:
                return pd.read_csv(csv_path)
            except UnicodeDecodeError:
                return pd.read_csv(csv_path, encoding="gbk")
            except Exception:
                return pd.DataFrame()

        self._case_tables = {
            "channels": _load_csv("channels"),
            "cases": _load_csv("cases"),
            "goals": _load_csv("goals"),
            "industry_brand": _load_csv("industry_brand"),
        }
        return self._case_tables

    def relation_detail_text(self) -> str:
        """Load docx relation details as plain text for reference."""
        doc_path = self.base_dir / self.files["relation_details"]
        if not doc_path.exists() or docx is None:
            return ""
        try:
            document = docx.Document(doc_path)
            parts = []
            for para in document.paragraphs:
                txt = para.text.strip()
                if txt:
                    parts.append(txt)
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception:
            return ""

    def alias_index(self) -> Dict[str, str]:
        """Build alias -> canonical mapping from case tables."""
        if self._alias_index:
            return self._alias_index
        tables = self.case_tables()
        alias_map: Dict[str, str] = {}

        # 渠道表：二级/三级作为 alias，一级作为 canonical 类别
        channel_df = tables.get("channels")
        if not channel_df.empty:
            for _, row in channel_df.iterrows():
                primary = str(row.get("一级", "")).strip() or str(row.get("一级1", "")).strip()
                secondary = str(row.get("二级1", "")).strip()
                tertiary = str(row.get("二级对应三级", "")).strip()
                canonical = primary or secondary
                for token in [primary, secondary]:
                    token = str(token).strip()
                    if token and token != "nan":
                        alias_map[token.lower()] = canonical
                if tertiary and tertiary != "nan":
                    for item in str(tertiary).split(","):
                        alias = item.strip()
                        if alias:
                            alias_map[alias.lower()] = canonical

        goal_df = tables.get("goals")
        if not goal_df.empty:
            for _, row in goal_df.iterrows():
                primary = str(row.get("一级分类", "")).strip()
                secondary = str(row.get("二级分类", "")).strip()
                canonical = secondary or primary
                for token in [primary, secondary]:
                    token = str(token).strip()
                    if token and token != "nan":
                        alias_map[token.lower()] = canonical

        industry_df = tables.get("industry_brand")
        if not industry_df.empty:
            for _, row in industry_df.iterrows():
                primary = str(row.get("一级行业分类", "")).strip()
                secondaries = str(row.get("二级行业分类", "")).split(",")
                for sec in secondaries:
                    alias = sec.strip()
                    if alias:
                        alias_map[alias.lower()] = primary
                if primary:
                    alias_map[primary.lower()] = primary

        # 案例库中的企业/品牌字段
        cases_df = tables.get("cases")
        if not cases_df.empty:
            for col in cases_df.columns:
                if "企业" in col or "品牌" in col:
                    for val in cases_df[col].dropna().tolist():
                        name = str(val).strip()
                        if name:
                            alias_map[name.lower()] = name

        self._alias_index = alias_map
        return self._alias_index

    def canonicalize(self, name: str) -> Tuple[str, Optional[str]]:
        """Return (canonical_name, source) if alias is known."""
        alias_map = self.alias_index()
        key = name.lower()
        if key in alias_map:
            return alias_map[key], "case_library"
        return name, None

    def schema_extension(self) -> SchemaExtension:
        """Build schema hints for Cypher/NER prompts."""
        if self._schema_extension:
            return self._schema_extension

        ext = SchemaExtension()
        tables = self.case_tables()

        # 渠道
        if not tables.get("channels", pd.DataFrame()).empty:
            ext.node_types.update({"ChannelCategory", "Channel"})
            ext.relationship_types.update({"HAS_CHANNEL", "BELONGS_TO_CHANNEL_CATEGORY"})
            ext.predicates.update({"uses_channel", "media_placement"})

        # 目标
        if not tables.get("goals", pd.DataFrame()).empty:
            ext.node_types.update({"PRGoal"})
            ext.relationship_types.update({"ALIGNS_WITH_GOAL"})
            ext.predicates.update({"targets"})

        # 行业-品牌
        if not tables.get("industry_brand", pd.DataFrame()).empty:
            ext.node_types.update({"Industry"})
            ext.relationship_types.update({"IN_INDUSTRY"})

        # 案例
        if not tables.get("cases", pd.DataFrame()).empty:
            ext.node_types.update({"PRCase", "CampaignAsset"})
            ext.relationship_types.update({"HAS_ASSET", "RELATED_TO_BRAND", "RELATED_TO_CHANNEL"})
            ext.predicates.update({"references_case"})

        # 关系详情补充
        detail_text = self.relation_detail_text()
        if detail_text:
            # 粗粒度提取可能的谓词关键词
            for token in ["合作", "投放", "联名", "冠名", "活动", "发布", "危机", "曝光", "流量"]:
                ext.predicates.add(token)

        self._schema_extension = ext
        return ext

    def to_metadata(self) -> Dict[str, str]:
        """Expose lightweight summary for logging/metadata."""
        tables = self.case_tables()
        return {
            "methodology_len": str(len(self.methodology_text())),
            "channels_rows": str(len(tables.get("channels", []))),
            "cases_rows": str(len(tables.get("cases", []))),
            "goals_rows": str(len(tables.get("goals", []))),
            "industry_rows": str(len(tables.get("industry_brand", []))),
        }

    def export_case_tables(self) -> Dict[str, List[Dict[str, str]]]:
        """Convert tables to records for downstream ingestion/Neo4j写入."""
        out: Dict[str, List[Dict[str, str]]] = {}
        for key, df in self.case_tables().items():
            if df.empty:
                out[key] = []
            else:
                out[key] = json.loads(df.fillna("").to_json(orient="records", force_ascii=False))
        return out


__all__ = ["ReferenceSources", "SchemaExtension"]
